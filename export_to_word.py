import pandas as pd
import ast
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches
import re
import regex
import string
from docxcompose.composer import Composer
from aife_time import now_and_choices

chinese_dun_ordinal = r"[零一二三四五六七八九十百]+、.*"
chinese_is_ordinal = r"[零一二三四五六七八九十百]+是.*"
chinese_bracket_ordinal = r"（[零一二三四五六七八九十百]+）.*"
arabic_dot_ordinal = r"(?:\d+\.)+(?!\d).*"
chinese_dun_ordinal_full_stop = r"[零一二三四五六七八九十百]+、.*?。"
chinese_is_ordinal_full_stop = r"[零一二三四五六七八九十百]+是.*?。"
chinese_bracket_ordinal_full_stop = r"（[零一二三四五六七八九十百]+）.*?。"
arabic_dot_ordinal_full_stop = r"(?:\d+\.)+(?!\d).*?。"

ordinal = f"{chinese_dun_ordinal}|{chinese_is_ordinal}|{chinese_bracket_ordinal}|{arabic_dot_ordinal}"
ordinal_full_stop = f"{chinese_dun_ordinal_full_stop}|{chinese_is_ordinal_full_stop}|{chinese_bracket_ordinal_full_stop}|{arabic_dot_ordinal_full_stop}"


def integrate_lines(body_content):
    lines = []
    for value in body_content.values():
        text_chunks = re.split(f"({ordinal_full_stop})", value)
        line = ""
        for i, text_chunk in enumerate(text_chunks):
            if re.match(f"^{ordinal_full_stop}$", text_chunk):
                if line:
                    lines.append(line.strip())
                line = f"**{text_chunk.strip()}**"
            elif re.match(f"^{ordinal}$", text_chunk):
                if line:
                    lines.append(line.strip())
                line = f"**{text_chunk.strip()}**"
            elif text_chunk.strip():
                if i == 0:
                    line = text_chunk.strip()
                else:
                    line += text_chunk.strip()
        lines.append(line.strip())
    return lines


def process_all_text_paragraphs(doc, *functions):
    paragraphs = doc.paragraphs
    for paragraph in paragraphs[8:]:
        if any(run.text.strip() for run in paragraph.runs):
            for function in functions:
                function(paragraph)


def copy_run_style(run, new_run):
    new_run.bold = run.bold
    new_run.font.name = run.font.name
    new_run.font.size = run.font.size


def replace_halfwidth_quotes_with_fullwidth(paragraph):
    pattern = re.compile(r'"')
    opening_quote = True
    new_runs = []
    for run in paragraph.runs:
        text = run.text
        if text:
            text_chunks = pattern.split(text)
            quotes = pattern.findall(text)
            for index, text_chunk in enumerate(text_chunks):
                if text_chunk:
                    new_runs.append((text_chunk, run))
                if index < len(quotes):
                    if opening_quote:
                        new_runs.append(("“", run))
                    else:
                        new_runs.append(("”", run))
                    opening_quote = not opening_quote
    for run in paragraph.runs:
        run.text = ""
    for text_chunk, run in new_runs:
        new_run = paragraph.add_run(text_chunk)
        copy_run_style(run, new_run)


def remove_special_symbols(paragraph):
    pattern = regex.compile(r"[^\p{Letter}\p{Number}\p{Han}\p{Punctuation}\p{Math_Symbol}\p{Currency_Symbol}\p{Z}]")
    new_runs = []
    for run in paragraph.runs:
        text = run.text
        if text:
            cleaned_text = pattern.sub("", text)
            new_runs.append((cleaned_text, run))
    for run in paragraph.runs:
        run.text = ""
    for cleaned_text, run in new_runs:
        new_run = paragraph.add_run(cleaned_text)
        copy_run_style(run, new_run)


def change_digits_letters_punctuation_to_times_new_roman(paragraph):
    pattern = re.compile(r"([0-9A-Za-z" + re.escape(string.punctuation) + r"]+)")
    new_runs = []
    for run in paragraph.runs:
        text = run.text
        if text:
            text_chunks = pattern.split(text)
            for index, text_chunk in enumerate(text_chunks):
                if text_chunk:
                    to_change = (index % 2 == 1)
                    new_runs.append((text_chunk, run, to_change))
    for run in paragraph.runs:
        run.text = ""
    for text_chunk, run, to_change in new_runs:
        new_run = paragraph.add_run(text_chunk)
        copy_run_style(run, new_run)
        if to_change:
            new_run.font.name = "Times New Roman"


def center_image_description_paragraphs(doc):
    paragraphs = doc.paragraphs
    for i, paragraph in enumerate(paragraphs):
        if any(run.text.strip() == "" and run.element.xpath(".//w:drawing") for run in paragraph.runs):
            prev_text_paragraph = None
            j = i - 1
            while j >= 0:
                if paragraphs[j].text.strip():
                    prev_text_paragraph = paragraphs[j]
                    break
                j -= 1

            next_text_paragraph = None
            j = i + 1
            while j < len(paragraphs):
                if paragraphs[j].text.strip():
                    next_text_paragraph = paragraphs[j]
                    break
                j += 1

            if prev_text_paragraph and "。" not in prev_text_paragraph.text:
                prev_text_paragraph.paragraph_format.first_line_indent = Pt(0)
                prev_text_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if next_text_paragraph and "。" not in next_text_paragraph.text:
                next_text_paragraph.paragraph_format.first_line_indent = Pt(0)
                next_text_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def export_search_results_to_word(csv_path):
    df = pd.read_csv(csv_path, encoding="utf-8")
    valid_mask = (df["heading_2"].notna() & df["source"].notna() & df["published_date"].notna() & df["body_content"].notna())
    doc = Document("ab_doc_temps/info_search_temp_start.docx")
    written_heading_1 = set()

    for index, row in df[valid_mask].iterrows():
        try:
            heading_1 = row["heading_1"] if pd.notna(row["heading_1"]) else None
            heading_2 = row["heading_2"]
            source = row["source"]
            published_date = row["published_date"]
            body_content = ast.literal_eval(row["body_content"])
            body_content = integrate_lines(body_content)

            if heading_1 and heading_1 not in written_heading_1:
                written_heading_1.add(heading_1)
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(heading_1)
                paragraph.style = doc.styles["Heading 1"]
                run.font.name = "楷体"
                run.font.size = Pt(22)
                run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            paragraph = doc.add_paragraph()
            run = paragraph.add_run(heading_2)
            paragraph.style = doc.styles["Heading 2"]
            run.font.name = "楷体"
            run.font.size = Pt(15)
            run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(6)

            paragraph = doc.add_paragraph()
            paragraph.style = "Normal"
            run_source = paragraph.add_run(source + " ")
            run_source.font.name = "宋体"
            run_source.font.size = Pt(12)
            run_date = paragraph.add_run(published_date)
            run_date.font.name = "Times New Roman"
            run_date.font.size = Pt(12)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for value in body_content:
                if value.startswith("temp-images"):
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(value, width=Inches(5.0))
                    paragraph.alignment = 1
                else:
                    paragraph = doc.add_paragraph()
                    paragraph.style = "Normal"
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraph.paragraph_format.first_line_indent = Pt(24)
                    paragraph.paragraph_format.line_spacing = 1.25

                    if "*" in value:
                        text_chunks = re.split(r"(\*\*.*?\*\*)", value)
                        for text_chunk in text_chunks:
                            if text_chunk.startswith("**") and text_chunk.endswith("**"):
                                text_chunk = text_chunk[2:-2]
                                run = paragraph.add_run(text_chunk)
                                run.bold = True
                            else:
                                text_chunk = text_chunk.replace("*", "")
                                run = paragraph.add_run(text_chunk)
                    else:
                        run = paragraph.add_run(value)
                    run.font.name = "宋体"
                    run.font.size = Pt(12)
        except Exception as e:
            print(f"Error processing value: {e}")

    process_all_text_paragraphs(doc, replace_halfwidth_quotes_with_fullwidth, remove_special_symbols, change_digits_letters_punctuation_to_times_new_roman)
    center_image_description_paragraphs(doc)

    doc_path = f"temp-data/{now_and_choices()}.docx"
    doc.save(doc_path)
    print(f"Document saved successfully: {doc_path}")
    return doc_path


def append_company_info_and_disclaimer(doc_path):
    doc = Document(doc_path)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    composer = Composer(doc)
    doc_to_append = Document("ab_doc_temps/info_search_temp_end.docx")
    composer.append(doc_to_append)
    composer.save(doc_path)
    print(f"Company info and disclaimer added to: {doc_path}")


if __name__ == "__main__":
    csv_path = r""
    doc_path = export_search_results_to_word(csv_path)
    append_company_info_and_disclaimer(doc_path)
    print(doc_path)
